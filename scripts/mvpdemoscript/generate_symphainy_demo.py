#!/usr/bin/env python3
import os
import json
import csv
import struct
import random
from faker import Faker
import pandas as pd
from pathlib import Path
from zipfile import ZipFile
from datetime import datetime, timedelta

fake = Faker()
BASE_DIR = Path("/home/founders/demoversion/symphainy_source/scripts/mvpdemoscript/demo_files")

SCENARIOS = ["Defense_TnE", "Underwriting_Insights", "Coexistence"]

# GCS Configuration
GCS_BUCKET_NAME = "symphainy-demo-files"  # Change this to your bucket name
GCS_PROJECT_ID = None  # Will auto-detect from credentials

def make_dirs(scenario):
    scenario_dir = BASE_DIR / f"SymphAIny_Demo_{scenario}"
    dirs = {
        "data": scenario_dir / "data",
        "metadata": scenario_dir / "metadata",
        "insights": scenario_dir / "insights",
        "workflow": scenario_dir / "workflow",
        "blueprint": scenario_dir / "blueprint",
        "visuals": scenario_dir / "visuals",
        "scripts": scenario_dir / "scripts",
    }
    for d in dirs.values():
        d.mkdir(parents=True, exist_ok=True)
    return dirs

# --------------------------
# Scenario A: Defense / T&E
# --------------------------
def generate_defense_tne(dirs):
    # 1. mission_plan.csv
    with open(dirs["data"] / "mission_plan.csv", "w", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=["mission_id", "start_time", "end_time", "location", "lead_officer"])
        writer.writeheader()
        for i in range(50):
            start = datetime.now() + timedelta(days=i)
            writer.writerow({
                "mission_id": f"M{i:03}",
                "start_time": start.isoformat(),
                "end_time": (start + timedelta(hours=2)).isoformat(),
                "location": fake.city(),
                "lead_officer": fake.name()
            })

    # 2. telemetry_raw.bin + copybook
    record_struct = "I7f"  # mission_id(int), timestamp(float), altitude(float), sensor1-5
    copybook_text = """      01  TELEMETRY-RECORD.
         05  MISSION-ID       PIC 9(4) COMP.
         05  TIMESTAMP        PIC 9(9)V9(6) COMP-3.
         05  ALTITUDE         PIC S9(5)V9(2) COMP-3.
         05  SENSOR1          PIC 9(5)V9(2).
         05  SENSOR2          PIC 9(5)V9(2).
         05  SENSOR3          PIC 9(5)V9(2).
         05  SENSOR4          PIC 9(5)V9(2).
         05  SENSOR5          PIC 9(5)V9(2)."""
    with open(dirs["metadata"] / "telemetry_copybook.cpy", "w") as f:
        f.write(copybook_text)

    with open(dirs["data"] / "telemetry_raw.bin", "wb") as f:
        for i in range(50):
            mission_id = i
            timestamp = float(i*1000)
            altitude = random.uniform(100, 1000)
            sensors = [random.uniform(0, 100) for _ in range(5)]
            packed = struct.pack("I7f", mission_id, timestamp, altitude, *sensors)
            f.write(packed)

    # 3. DOCX incident logs
    from docx import Document
    doc = Document()
    for i in range(3):
        doc.add_heading(f"Incident {i+1}", level=2)
        doc.add_paragraph(f"Mission {i}, weather: {fake.word()}, summary: {fake.sentence()}")
    doc.save(dirs["data"] / "test_incident_reports.docx")

    # 4. Insights JSON
    insights = {"coverage_gaps_pct": 12, "anomalies_detected": 3}
    with open(dirs["insights"] / "ai_insights.json", "w") as f:
        json.dump(insights, f, indent=2)

    # 5. Workflow & blueprint placeholders
    dirs["workflow"].joinpath("workflow_mermaid.mmd").write_text("graph TD; A-->B; B-->C")
    dirs["workflow"].joinpath("workflow.bpmn").write_text("<bpmn>...</bpmn>")
    dirs["blueprint"].joinpath("coexistence_blueprint.png").write_text("PLACEHOLDER PNG")

# --------------------------
# Scenario B: Underwriting Insights
# --------------------------
def generate_underwriting_insights(dirs):
    n_rows = 10000
    # Binary policy_master.dat
    record_struct = "10s10sIf?"  # policy_id, dob, face, status, smoker_flag
    with open(dirs["metadata"] / "copybook.cpy", "w") as f:
        f.write("Copybook placeholder for 80-byte fixed-width record")

    with open(dirs["data"] / "policy_master.dat", "wb") as f:
        for _ in range(n_rows):
            policy_id = fake.bothify(text='??#####').ljust(10)[:10].encode("utf-8")
            dob = fake.date_of_birth().strftime("%Y-%m-%d").encode("utf-8")
            face = random.uniform(10000, 500000)
            status = random.randint(0,1)
            smoker = random.choice([0,1])
            packed = struct.pack("10s10sIf?", policy_id, dob, int(face), float(status), bool(smoker))
            f.write(packed)

    # claims.csv
    df = pd.DataFrame({
        "policy_id":[fake.bothify(text='??#####') for _ in range(5000)],
        "claim_amount":[random.uniform(500, 50000) for _ in range(5000)],
        "claim_date":[fake.date_this_decade() for _ in range(5000)]
    })
    df.to_csv(dirs["data"] / "claims.csv", index=False)

    # reinsurance.xlsx
    df.to_excel(dirs["data"] / "reinsurance.xlsx", index=False)

    # underwriting notes PDF
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    for _ in range(50):
        pdf.multi_cell(w=pdf.epw, h=5, text=fake.paragraph())
    pdf.output(str(dirs["data"] / "underwriting_notes.pdf"))

    # AI insights & workflow/blueprint placeholders
    insights = {"mortality_trends_pct": 15, "reserving_patterns": "synthetic_summary"}
    with open(dirs["insights"] / "ai_insights.json", "w") as f:
        json.dump(insights, f, indent=2)
    dirs["workflow"].joinpath("workflow_mermaid.mmd").write_text("graph TD; X-->Y; Y-->Z")
    dirs["workflow"].joinpath("workflow.bpmn").write_text("<bpmn>...</bpmn>")
    dirs["blueprint"].joinpath("coexistence_blueprint.png").write_text("PLACEHOLDER PNG")

# --------------------------
# Scenario C: Pre-Migration Coexistence
# --------------------------
def generate_pre_migration_coexistence(dirs):
    # legacy_policy_export.csv
    df = pd.DataFrame({
        "pol_num":[fake.bothify(text='??#####') for _ in range(50)],
        "holder_name":[fake.name() for _ in range(50)],
        "dob":[fake.date_of_birth() for _ in range(50)],
        "premium":[random.uniform(100,1000) for _ in range(50)]
    })
    df.to_csv(dirs["data"] / "legacy_policy_export.csv", index=False)

    # target schema JSON
    schema = {"policy_id":"string","name":"string","dob":"date","premium":"float"}
    with open(dirs["data"] / "target_schema.json", "w") as f:
        json.dump(schema, f, indent=2)

    # alignment map
    alignment = {"pol_num":"policy_id","holder_name":"name"}
    with open(dirs["metadata"] / "alignment_map.json", "w") as f:
        json.dump(alignment, f, indent=2)

    # AI insights & workflow/blueprint placeholders
    insights = {"coverage_pct": 95, "type_mismatches": 2}
    with open(dirs["insights"] / "ai_insights.json", "w") as f:
        json.dump(insights, f, indent=2)
    dirs["workflow"].joinpath("workflow_mermaid.mmd").write_text("graph TD; L-->M; M-->N")
    dirs["workflow"].joinpath("workflow.bpmn").write_text("<bpmn>...</bpmn>")
    dirs["blueprint"].joinpath("coexistence_blueprint.png").write_text("PLACEHOLDER PNG")

# --------------------------
# ZIP Packaging
# --------------------------
def zip_scenario(scenario):
    scenario_dir = BASE_DIR / f"SymphAIny_Demo_{scenario}"
    zip_name = f"{scenario_dir}.zip"
    with ZipFile(zip_name, 'w') as zipf:
        for folder, _, files in os.walk(scenario_dir):
            for file in files:
                file_path = os.path.join(folder, file)
                zipf.write(file_path, os.path.relpath(file_path, scenario_dir))
    print(f"  📦 Created ZIP: {os.path.basename(zip_name)}")
    return zip_name

# --------------------------
# GCS Upload with Signed URLs
# --------------------------
def upload_to_gcs(zip_path, bucket_name, expire_hours=24):
    """
    Upload zip file to Google Cloud Storage and generate a signed URL.
    
    Args:
        zip_path: Path to the zip file
        bucket_name: Name of the GCS bucket
        expire_hours: How long the signed URL should be valid (default 24 hours)
    
    Returns:
        tuple: (public_url, signed_url, blob_name)
    """
    try:
        from google.cloud import storage
        
        # Initialize GCS client
        storage_client = storage.Client(project=GCS_PROJECT_ID)
        bucket = storage_client.bucket(bucket_name)
        
        # Upload file
        blob_name = f"demo_files/{os.path.basename(zip_path)}"
        blob = bucket.blob(blob_name)
        
        print(f"  📤 Uploading {os.path.basename(zip_path)} to GCS...")
        blob.upload_from_filename(str(zip_path))
        
        # Generate signed URL (valid for specified hours)
        signed_url = blob.generate_signed_url(
            version="v4",
            expiration=timedelta(hours=expire_hours),
            method="GET"
        )
        
        # Public URL (only works if bucket is public)
        public_url = blob.public_url
        
        print(f"  ✅ Upload successful!")
        return public_url, signed_url, blob_name
        
    except ImportError:
        print("  ❌ ERROR: google-cloud-storage not installed")
        print("  Run: pip install google-cloud-storage")
        return None, None, None
    except Exception as e:
        print(f"  ❌ ERROR uploading to GCS: {e}")
        return None, None, None

# --------------------------
# Main
# --------------------------
if __name__ == "__main__":
    print("=" * 80)
    print("🎯 SymphAIny Demo Data Generator")
    print("=" * 80)
    print(f"📁 Base directory: {BASE_DIR}")
    print(f"☁️  GCS Bucket: {GCS_BUCKET_NAME}")
    print(f"📦 Scenarios: {', '.join(SCENARIOS)}")
    print("=" * 80)
    print()
    
    download_urls = {}
    
    for scenario in SCENARIOS:
        print(f"\n🔨 Generating {scenario}...")
        dirs = make_dirs(scenario)
        
        if scenario == "Defense_TnE":
            generate_defense_tne(dirs)
        elif scenario == "Underwriting_Insights":
            generate_underwriting_insights(dirs)
        elif scenario == "Coexistence":
            generate_pre_migration_coexistence(dirs)
        
        # Create zip file
        zip_path = zip_scenario(scenario)
        print(f"✅ {scenario} demo data generated!")
        
        # Upload to GCS
        public_url, signed_url, blob_name = upload_to_gcs(zip_path, GCS_BUCKET_NAME, expire_hours=168)  # 7 days
        
        if signed_url:
            download_urls[scenario] = {
                "signed_url": signed_url,
                "blob_name": blob_name,
                "file_size": os.path.getsize(zip_path) / (1024 * 1024)  # MB
            }
    
    # Print summary with download instructions
    print("\n" + "=" * 80)
    print("🎉 ALL DEMO FILES GENERATED & UPLOADED!")
    print("=" * 80)
    print()
    
    if download_urls:
        print("📥 DOWNLOAD LINKS (Valid for 7 days):")
        print()
        for scenario, info in download_urls.items():
            print(f"  📦 {scenario}:")
            print(f"     Size: {info['file_size']:.2f} MB")
            print(f"     GCS Path: gs://{GCS_BUCKET_NAME}/{info['blob_name']}")
            print(f"     Download URL:")
            print(f"     {info['signed_url']}")
            print()
        
        print("=" * 80)
        print("💡 HOW TO USE THESE FILES:")
        print("=" * 80)
        print()
        print("Option 1: Direct Download")
        print("  → Click any of the download URLs above")
        print("  → Files will download to your local machine")
        print("  → Upload them via the SymphAIny web interface")
        print()
        print("Option 2: Google Cloud Console")
        print("  → Go to: https://console.cloud.google.com/storage/browser")
        print(f"  → Navigate to bucket: {GCS_BUCKET_NAME}")
        print("  → Find files in the 'demo_files/' folder")
        print("  → Click ⋮ → Download")
        print()
        print("Option 3: gsutil Command Line")
        print(f"  → gsutil cp gs://{GCS_BUCKET_NAME}/demo_files/*.zip .")
        print()
    else:
        print("⚠️  No files were uploaded to GCS. Check errors above.")
        print(f"📁 Local files are available in: {BASE_DIR}")
    
    print("=" * 80)

