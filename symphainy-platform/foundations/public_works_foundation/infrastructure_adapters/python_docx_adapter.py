#!/usr/bin/env python3
"""
Python-DOCX Adapter - Raw Technology Client

Raw python-docx client wrapper for Word document operations.
This is Layer 1 of the 5-layer infrastructure architecture.

WHAT (Infrastructure Role): I provide raw python-docx operations
HOW (Infrastructure Implementation): I use the python-docx library with no business logic
"""

from typing import Dict, Any, Optional, List
from datetime import datetime
import logging
import io

try:
    from docx import Document
except ImportError:
    class Document:
        def __init__(self, file_path): pass
        @property
        def paragraphs(self): return []
        @property
        def core_properties(self): return None

logger = logging.getLogger(__name__)

class PythonDocxAdapter:
    """
    Raw python-docx client wrapper - no business logic.
    
    This adapter provides direct access to python-docx operations for
    Word document processing.
    """
    
    def __init__(self):
        """Initialize Python-DOCX Adapter."""
        logger.info("✅ Python-DOCX Adapter initialized")
    
    async def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract raw text from Word document.
        
        Args:
            file_path: Path to the Word document.
            
        Returns:
            Dict[str, Any]: A dictionary containing the extracted text and metadata.
        """
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                "success": True,
                "text": text.strip(),
                "paragraph_count": len(doc.paragraphs),
                "timestamp": datetime.utcnow().isoformat()
            }
        except Exception as e:
            logger.error(f"❌ Python-DOCX text extraction failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "paragraph_count": 0,
                "timestamp": datetime.utcnow().isoformat()
            }
    
    async def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from Word document.
        
        Args:
            file_path: Path to the Word document.
            
        Returns:
            Dict[str, Any]: A dictionary containing the document metadata.
        """
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "created": core_props.created,
                "modified": core_props.modified,
                "subject": core_props.subject,
                "keywords": core_props.keywords
            }
            
            return {
                "success": True,
                "metadata": metadata,
                "timestamp": datetime.utcnow().isoformat()
            }
        except Exception as e:
            logger.error(f"❌ Python-DOCX metadata extraction failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "metadata": {},
                "timestamp": datetime.utcnow().isoformat()
            }
    
    async def extract_text_from_bytes(self, file_data: bytes) -> Dict[str, Any]:
        """
        Extract raw text from Word document bytes.
        
        Args:
            file_data: Word document content as bytes
            
        Returns:
            Dict[str, Any]: A dictionary containing the extracted text and metadata.
        """
        try:
            doc = Document(io.BytesIO(file_data))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                "success": True,
                "text": text.strip(),
                "paragraph_count": len(doc.paragraphs),
                "timestamp": datetime.utcnow().isoformat()
            }
        except Exception as e:
            logger.error(f"❌ Python-DOCX text extraction from bytes failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "paragraph_count": 0,
                "timestamp": datetime.utcnow().isoformat()
            }
    
    async def extract_metadata_from_bytes(self, file_data: bytes) -> Dict[str, Any]:
        """
        Extract metadata from Word document bytes.
        
        Args:
            file_data: Word document content as bytes
            
        Returns:
            Dict[str, Any]: A dictionary containing the document metadata.
        """
        try:
            doc = Document(io.BytesIO(file_data))
            core_props = doc.core_properties
            
            metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "created": core_props.created,
                "modified": core_props.modified,
                "subject": core_props.subject,
                "keywords": core_props.keywords
            }
            
            return {
                "success": True,
                "metadata": metadata,
                "timestamp": datetime.utcnow().isoformat()
            }
        except Exception as e:
            logger.error(f"❌ Python-DOCX metadata extraction from bytes failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "metadata": {},
                "timestamp": datetime.utcnow().isoformat()
            }
    
    async def get_status(self) -> Dict[str, Any]:
        """Get the status of the adapter."""
        return {
            "adapter_name": "PythonDocxAdapter",
            "status": "ready",
            "library_version": getattr(Document, '__version__', 'unknown'),
            "timestamp": datetime.utcnow().isoformat()
        }






