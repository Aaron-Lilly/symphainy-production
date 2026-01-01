#!/usr/bin/env python3
"""
Test File Helpers for File Parser Functional Tests

Creates real test files (Excel, Word, PDF, binary, copybook) for functional testing.
"""

import io
import tempfile
import os
import json
import csv
from typing import Dict, Any, Optional, Tuple
from pathlib import Path


def create_test_json_file(data: Optional[Dict[str, Any]] = None) -> bytes:
    """
    Create a test JSON file with sample data.
    
    Args:
        data: Optional dictionary to serialize. If None, uses default test data.
    
    Returns:
        JSON file content as bytes
    """
    if data is None:
        data = {
            "name": "John Doe",
            "age": 30,
            "email": "john@example.com",
            "balance": 1000.50
        }
    
    json_str = json.dumps(data, indent=2)
    return json_str.encode('utf-8')


def create_test_csv_file(data: Optional[list] = None) -> bytes:
    """
    Create a test CSV file with sample data.
    
    Args:
        data: Optional list of dictionaries to serialize. If None, uses default test data.
    
    Returns:
        CSV file content as bytes
    """
    if data is None:
        data = [
            {"Name": "John", "Age": 30, "City": "New York"},
            {"Name": "Jane", "Age": 25, "City": "Boston"},
            {"Name": "Bob", "Age": 35, "City": "Chicago"}
        ]
    
    output = io.StringIO()
    if data:
        writer = csv.DictWriter(output, fieldnames=data[0].keys())
        writer.writeheader()
        writer.writerows(data)
    
    return output.getvalue().encode('utf-8')


def create_test_text_file(content: Optional[str] = None) -> bytes:
    """
    Create a test text file with sample content.
    
    Args:
        content: Optional text content. If None, uses default test content.
    
    Returns:
        Text file content as bytes
    """
    if content is None:
        content = """Test Text File
This is a test text file for validation and export testing.
It contains multiple lines of text.
Line 4
Line 5"""
    
    return content.encode('utf-8')


def create_test_excel_file() -> Tuple[bytes, str]:
    """
    Create a test Excel file with sample data.
    
    Returns:
        Tuple of (file_bytes, filename)
    """
    try:
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.utils.dataframe import dataframe_to_rows
        
        # Create test data
        data = {
            "Name": ["Alice", "Bob", "Charlie", "Diana"],
            "Age": [25, 30, 35, 28],
            "City": ["New York", "London", "Tokyo", "Paris"],
            "Salary": [50000, 60000, 70000, 55000]
        }
        df = pd.DataFrame(data)
        
        # Create Excel file in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Employees', index=False)
        
        output.seek(0)
        return output.getvalue(), "test_employees.xlsx"
        
    except ImportError:
        # Fallback: Create minimal Excel file using openpyxl directly
        try:
            from openpyxl import Workbook
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Employees"
            
            # Headers
            ws.append(["Name", "Age", "City", "Salary"])
            
            # Data rows
            ws.append(["Alice", 25, "New York", 50000])
            ws.append(["Bob", 30, "London", 60000])
            ws.append(["Charlie", 35, "Tokyo", 70000])
            ws.append(["Diana", 28, "Paris", 55000])
            
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            return output.getvalue(), "test_employees.xlsx"
            
        except ImportError:
            raise ImportError("openpyxl or pandas required for Excel test file creation")


def create_test_word_document() -> Tuple[bytes, str]:
    """
    Create a test Word document with sample content.
    
    Returns:
        Tuple of (file_bytes, filename)
    """
    try:
        from docx import Document
        
        doc = Document()
        
        # Add title
        doc.add_heading('Test Document', 0)
        
        # Add paragraphs
        doc.add_paragraph('This is a test document for File Parser functional testing.')
        doc.add_paragraph('It contains multiple paragraphs and a table.')
        
        # Add table
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Item'
        header_cells[1].text = 'Quantity'
        header_cells[2].text = 'Price'
        
        # Data rows
        data_rows = [
            ['Apple', '10', '$1.00'],
            ['Banana', '5', '$0.50']
        ]
        for row_data in data_rows:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row_data):
                row_cells[i].text = cell_value
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue(), "test_document.docx"
        
    except ImportError:
        raise ImportError("python-docx required for Word document test file creation")


def create_test_pdf_file() -> Tuple[bytes, str]:
    """
    Create a test PDF file with sample content.
    
    Returns:
        Tuple of (file_bytes, filename)
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        output = io.BytesIO()
        c = canvas.Canvas(output, pagesize=letter)
        
        # Add content
        c.drawString(100, 750, "Test PDF Document")
        c.drawString(100, 730, "This is a test PDF for File Parser functional testing.")
        c.drawString(100, 710, "It contains text content that should be extractable.")
        
        # Add a table-like structure
        c.drawString(100, 680, "Item | Quantity | Price")
        c.drawString(100, 660, "Apple | 10 | $1.00")
        c.drawString(100, 640, "Banana | 5 | $0.50")
        
        c.save()
        output.seek(0)
        return output.getvalue(), "test_document.pdf"
        
    except ImportError:
        raise ImportError("reportlab required for PDF test file creation")


def create_test_copybook_file() -> Tuple[bytes, str]:
    """
    Create a test COBOL copybook file with field definitions.
    
    Returns:
        Tuple of (file_bytes, filename)
    """
    copybook_content = """       01  CUSTOMER-RECORD.
           05  CUSTOMER-ID          PIC X(10).
           05  CUSTOMER-NAME        PIC X(30).
           05  CUSTOMER-AGE         PIC 9(3).
           05  CUSTOMER-SALARY      PIC 9(8)V99 COMP-3.
           05  CUSTOMER-STATUS     PIC X(1).
           05  CUSTOMER-DATE       PIC X(8).
"""
    return copybook_content.encode('utf-8'), "test_copybook.cpy"


def create_test_binary_file(copybook_path: Optional[str] = None) -> Tuple[bytes, str]:
    """
    Create a test binary file with sample mainframe data.
    
    Uses the copybook structure to create properly formatted binary records.
    
    Args:
        copybook_path: Optional path to copybook file for field definitions
    
    Returns:
        Tuple of (file_bytes, filename)
    """
    # Based on the copybook structure:
    # CUSTOMER-ID: 10 bytes (X(10))
    # CUSTOMER-NAME: 30 bytes (X(30))
    # CUSTOMER-AGE: 3 bytes (9(3))
    # CUSTOMER-SALARY: 5 bytes (COMP-3, 8 digits + 2 decimal = 5 bytes)
    # CUSTOMER-STATUS: 1 byte (X(1))
    # CUSTOMER-DATE: 8 bytes (X(8))
    # Total record length: 10 + 30 + 3 + 5 + 1 + 8 = 57 bytes
    
    records = []
    
    # Record 1
    record1 = bytearray()
    record1.extend(b"CUST001    ")  # CUSTOMER-ID (10 bytes)
    record1.extend(b"Alice Smith                    ")  # CUSTOMER-NAME (30 bytes)
    record1.extend(b"025")  # CUSTOMER-AGE (3 bytes)
    record1.extend(_pack_comp3(50000.00))  # CUSTOMER-SALARY (5 bytes COMP-3)
    record1.extend(b"A")  # CUSTOMER-STATUS (1 byte)
    record1.extend(b"20240101")  # CUSTOMER-DATE (8 bytes)
    records.append(bytes(record1))
    
    # Record 2
    record2 = bytearray()
    record2.extend(b"CUST002    ")  # CUSTOMER-ID (10 bytes)
    record2.extend(b"Bob Johnson                    ")  # CUSTOMER-NAME (30 bytes)
    record2.extend(b"030")  # CUSTOMER-AGE (3 bytes)
    record2.extend(_pack_comp3(60000.00))  # CUSTOMER-SALARY (5 bytes COMP-3)
    record2.extend(b"A")  # CUSTOMER-STATUS (1 byte)
    record2.extend(b"20240102")  # CUSTOMER-DATE (8 bytes)
    records.append(bytes(record2))
    
    # Record 3
    record3 = bytearray()
    record3.extend(b"CUST003    ")  # CUSTOMER-ID (10 bytes)
    record3.extend(b"Charlie Brown                  ")  # CUSTOMER-NAME (30 bytes)
    record3.extend(b"035")  # CUSTOMER-AGE (3 bytes)
    record3.extend(_pack_comp3(70000.00))  # CUSTOMER-SALARY (5 bytes COMP-3)
    record3.extend(b"I")  # CUSTOMER-STATUS (1 byte)
    record3.extend(b"20240103")  # CUSTOMER-DATE (8 bytes)
    records.append(bytes(record3))
    
    # Combine all records
    binary_data = b''.join(records)
    return binary_data, "test_binary.bin"


def _pack_comp3(value: float) -> bytes:
    """
    Pack a numeric value into COMP-3 (packed decimal) format.
    
    COMP-3 format: Each digit is stored as a nibble (4 bits), with the sign in the last nibble.
    For 8 digits + 2 decimal places = 10 digits total = 5 bytes (2 digits per byte + sign)
    
    Args:
        value: Numeric value to pack
    
    Returns:
        Packed bytes in COMP-3 format
    """
    # Convert to integer (multiply by 100 for 2 decimal places)
    int_value = int(value * 100)
    
    # Get absolute value
    abs_value = abs(int_value)
    
    # Convert to string of digits
    digits = str(abs_value).zfill(10)  # 10 digits total
    
    # Pack into bytes (2 digits per byte)
    packed = bytearray(5)
    for i in range(5):
        if i < 4:
            # Pack two digits
            byte_val = (int(digits[i * 2]) << 4) | int(digits[i * 2 + 1])
            packed[i] = byte_val
        else:
            # Last byte: last digit + sign
            last_digit = int(digits[8])
            sign = 0x0C if int_value >= 0 else 0x0D  # C for positive, D for negative
            packed[i] = (last_digit << 4) | sign
    
    return bytes(packed)


def create_test_unsupported_file() -> Tuple[bytes, str]:
    """
    Create a test file with unsupported format.
    
    Returns:
        Tuple of (file_bytes, filename)
    """
    content = b"This is a test file with an unsupported format."
    return content, "test_file.xyz"


def create_test_image_file() -> Tuple[bytes, str]:
    """
    Create a test image file with text (for OCR testing).
    
    Returns:
        Tuple of (file_bytes, filename)
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
        
        # Create a simple image with text
        img = Image.new('RGB', (400, 200), color='white')
        draw = ImageDraw.Draw(img)
        
        # Draw text
        try:
            # Try to use a default font
            font = ImageFont.load_default()
        except:
            font = None
        
        draw.text((10, 10), "Test Image for OCR", fill='black', font=font)
        draw.text((10, 40), "This text should be extractable", fill='black', font=font)
        draw.text((10, 70), "via OCR functionality", fill='black', font=font)
        
        # Save to bytes
        output = io.BytesIO()
        img.save(output, format='PNG')
        output.seek(0)
        return output.getvalue(), "test_image.png"
        
    except ImportError:
        raise ImportError("Pillow (PIL) required for image test file creation")


